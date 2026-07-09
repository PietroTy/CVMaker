"""
modules/exporter.py — Exportação do Currículo Adaptado (PDF, DOCX, TXT)
"""

import io
from datetime import datetime


def exportar(secoes_geradas: dict, formato: str, nome_candidato: str, cargo_alvo: str) -> tuple[bytes, str, str]:
    """
    Retorna (bytes, nome_arquivo, mime_type)
    secoes_geradas: {secao_id: {"titulo": str, "texto": str}}
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    nome_base = f"curriculo_{nome_candidato.replace(' ', '_')}_{ts}" if nome_candidato else f"curriculo_{ts}"

    if formato == "txt":
        return _gerar_txt(secoes_geradas, nome_candidato, cargo_alvo), f"{nome_base}.txt", "text/plain"
    elif formato == "pdf":
        return _gerar_pdf(secoes_geradas, nome_candidato, cargo_alvo), f"{nome_base}.pdf", "application/pdf"
    elif formato == "docx":
        return _gerar_docx(secoes_geradas, nome_candidato, cargo_alvo), f"{nome_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        return _gerar_txt(secoes_geradas, nome_candidato, cargo_alvo), f"{nome_base}.txt", "text/plain"


# ─── TXT ──────────────────────────────────────────────────────────
def _gerar_txt(secoes: dict, nome: str, cargo: str) -> bytes:
    linhas = []
    if nome:
        linhas.append(nome.upper())
        linhas.append("=" * len(nome))
    if cargo:
        linhas.append(f"Candidatura: {cargo}")
    linhas.append(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    linhas.append("")

    for sid, dados in secoes.items():
        titulo = dados.get("titulo", "")
        texto  = dados.get("texto", "")
        if not texto.strip():
            continue
        linhas.append(f"\n{'─' * 60}")
        linhas.append(titulo.upper())
        linhas.append('─' * 60)
        linhas.append(texto)

    linhas.append(f"\n{'─' * 60}")
    linhas.append("Gerado com CV Adapter • Powered by Claude AI")
    return "\n".join(linhas).encode("utf-8")


# ─── PDF ──────────────────────────────────────────────────────────
def _gerar_pdf(secoes: dict, nome: str, cargo: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib import colors
    except ImportError:
        raise ImportError("reportlab não instalado. Execute: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    cor_primaria = colors.HexColor("#1e3a5f")
    cor_acento   = colors.HexColor("#2563eb")

    styles.add(ParagraphStyle("NomeCandidato",
        fontName="Helvetica-Bold", fontSize=22, leading=26,
        textColor=cor_primaria, alignment=TA_CENTER, spaceAfter=4))
    styles.add(ParagraphStyle("CargoAlvo",
        fontName="Helvetica", fontSize=12, leading=14,
        textColor=cor_acento, alignment=TA_CENTER, spaceAfter=16))
    styles.add(ParagraphStyle("SecaoTitulo",
        fontName="Helvetica-Bold", fontSize=11, leading=14,
        textColor=cor_primaria, spaceBefore=14, spaceAfter=6,
        borderPad=(0,0,2,0)))
    styles.add(ParagraphStyle("CorpoTexto",
        fontName="Helvetica", fontSize=10, leading=14,
        textColor=colors.HexColor("#1a1a2e"), spaceAfter=4))
    styles.add(ParagraphStyle("Rodape",
        fontName="Helvetica", fontSize=8, leading=10,
        textColor=colors.gray, alignment=TA_CENTER))

    story = []

    # Cabeçalho
    if nome:
        story.append(Paragraph(nome, styles["NomeCandidato"]))
    if cargo:
        story.append(Paragraph(f"Candidatura: {cargo}", styles["CargoAlvo"]))
    story.append(HRFlowable(width="100%", thickness=2, color=cor_primaria, spaceAfter=8))

    # Seções
    for sid, dados in secoes.items():
        titulo = dados.get("titulo", "")
        texto  = dados.get("texto", "").strip()
        if not texto:
            continue

        story.append(Paragraph(titulo.upper(), styles["SecaoTitulo"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=cor_acento, spaceAfter=4))

        for linha in texto.split("\n"):
            linha = linha.strip()
            if linha:
                # Converte markdown básico
                linha = linha.replace("**", "<b>").replace("__", "<b>")
                if linha.startswith("- ") or linha.startswith("• "):
                    linha = "• " + linha[2:]
                story.append(Paragraph(linha, styles["CorpoTexto"]))

        story.append(Spacer(1, 6))

    # Rodapé
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceBefore=10, spaceAfter=4))
    story.append(Paragraph(f"Gerado com CV Adapter • {datetime.now().strftime('%d/%m/%Y')}", styles["Rodape"]))

    doc.build(story)
    return buf.getvalue()


# ─── DOCX ─────────────────────────────────────────────────────────
def _gerar_docx(secoes: dict, nome: str, cargo: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx não instalado. Execute: pip install python-docx")

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    cor_primaria = RGBColor(0x1e, 0x3a, 0x5f)
    cor_acento   = RGBColor(0x25, 0x63, 0xeb)

    # Nome
    if nome:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nome.upper())
        run.bold      = True
        run.font.size = Pt(22)
        run.font.color.rgb = cor_primaria

    # Cargo alvo
    if cargo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Candidatura: {cargo}")
        run.font.size = Pt(12)
        run.font.color.rgb = cor_acento

    doc.add_paragraph()

    # Seções
    for sid, dados in secoes.items():
        titulo = dados.get("titulo", "")
        texto  = dados.get("texto", "").strip()
        if not texto:
            continue

        # Título da seção
        h = doc.add_paragraph()
        run = h.add_run(titulo.upper())
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = cor_primaria

        # Linha separadora via underline no parágrafo
        doc.add_paragraph("─" * 60)

        # Corpo
        for linha in texto.split("\n"):
            linha = linha.strip()
            if linha:
                p    = doc.add_paragraph()
                run  = p.add_run(linha)
                run.font.size = Pt(10)

        doc.add_paragraph()

    # Rodapé
    doc.add_paragraph("─" * 60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Gerado com CV Adapter • {datetime.now().strftime('%d/%m/%Y')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
